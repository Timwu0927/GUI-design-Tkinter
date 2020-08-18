import tkinter as tk
import outlook
from outlook import outlookrun
import docx
from docx.shared import RGBColor
from docx.shared import Cm, Pt
import math
from functools import partial
import os
from docx2pdf import convert

import Interface_FTP
import datetime
import Interface_SQL
import time

#變數宣告dictionary
filename=locals()
chkValue = locals()
chkExample = locals()
chkValue2=locals()
chkExample2=locals()

def onOK():
    # 取得輸入文字
    try:
        all_data=outlook.outlookrun(name.get())
        data_counter=len(all_data)
        if(data_counter==0):  #收件匣中無信件時通知
            #print("寄件匣中沒有與PSA[公告]相關的信件，請確認後再執行")
            label_alert1 = tk.Label(window, text = "寄件匣中沒有與PSA[公告]相關的信件，請確認後再執行!!!")
            label_alert1.pack()
        else:
            for i in range(data_counter):
                chkValue['chkValue' + str(i) ] =  tk.BooleanVar()
                chkExample['chkExample' + str(i) ] = tk.Checkbutton(window,variable=chkValue.get('chkValue' + str(i)),text=all_data[i]['title']) 
                # chkValue.get('chkValue' + str(i))
                chkExample.get('chkExample' + str(i)).pack()
            
            button2 = tk.Button(window, text = "轉檔",command=partial(allfunc,data_counter)) #設置搜尋按鈕
            button2.pack()
    except:
        label_alert = tk.Label(window, text = "郵件地址輸入錯誤，請重新輸入!!!") #郵件地址輸入錯誤時通知
        label_alert.pack()

def allfunc(howlong):
    all_data=outlook.outlookrun(name.get())
    final=[]
    word=[] 
    for i in range(howlong):
        #print((chkExample.get('chkValue' + str(i))).get())
        if((chkExample.get('chkValue' + str(i))).get() == True):
            #print(1)
            #print(chkExample['chkExample'+str(i)]['text'])
            final.append({'title':chkExample['chkExample' + str(i)]['text']})
        
    for j in range(len(final)):
        for x in range(len(all_data)):
            if(final[j]['title']==all_data[x]['title']):
                word.append({'title':final[j]['title'],'content':all_data[x]['content']})


    temp_counter=len(word)
    for y in range(temp_counter):
        doc = docx.Document()                       #產生Word黨
        doc.add_paragraph(str(word[y]['content']))  #產生內文
        for paragraph in doc.paragraphs:            #設定文字型態、顏色和大小
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = '微軟正黑體'
                run.font.size = Pt(12)
                #print(run.text)
        if('圖' in word[y]['content']):             #判斷是否有圖片在信件中
            label_view = tk.Label(window, text = str(word[y]['title']+'\n 此標題內文中有圖片，不提供轉檔'))
            label_view.pack()
        else:   
            if('/' or '*' or '?' or '>' or '<' or ':' or '|' in (word[y]['title'])):   
                #print(all_data[j]['title'])
                #例外處理(將Word不接受的標題命名符號移除)
                word[y]['title']= word[y]['title'].replace('/','')
                word[y]['title']= word[y]['title'].replace('*','')
                word[y]['title']= word[y]['title'].replace('?','')
                word[y]['title']= word[y]['title'].replace('>','')
                word[y]['title']= word[y]['title'].replace('<','')
                word[y]['title']= word[y]['title'].replace(':','')
                word[y]['title']= word[y]['title'].replace('|','')
                doc.save(str(word[y]['title'])+'.docx') #存Word檔
                convert(str(word[y]['title']+'.docx'))  #轉成PDF檔
                #print(all_data[j]['title'])
                filename['filename'+str(y)]=os.path.abspath(str(word[y]['title']+'.docx'))

            else:
                doc.save(str(word[y]['title'])+'.docx')
                convert(str(word[y]['title']+'.docx'))
                filename['filename'+str(y)]=os.path.abspath(str(word[y]['title']+'.docx'))

#設置視窗大小與名稱   
window = tk.Tk()
window.title('Outlook 信件系統')
window.geometry("400x500+450+150")

#設定icon
window.iconphoto(True,tk.PhotoImage(file='/Users/tim.wu/Desktop/outlook怪/outlook.png'))

# 標示文字
label = tk.Label(window, text = '輸入NEC Mail地址')
label.pack()

# 輸入欄位
name = tk.Entry(window,     # 輸入欄位所在視窗
            width = 20)     # 輸入欄位的寬度
name.pack()

# labe2 = tk.Label(window, text = '輸入想要搜尋的關鍵字  ex: 收件匣')
# labe2.pack()

# position = tk.Entry(window,     # 輸入欄位所在視窗
#                  width = 20)    # 輸入欄位的寬度
# position.pack()

button = tk.Button(window, text = "搜尋", command = onOK) #設置搜尋按鈕
button.pack()

labe3 = tk.Label(window, text = '選擇需要放入Dashboard的相關公告')
labe3.pack()    

label4 = tk.Label(window, text = '註:若信件中有圖片，請手動更新後放至與程式相同目錄，再點選上傳!')
label4.pack(side='bottom')

upload_buttons=locals()
upload_buttons_value=locals()
uplaod_path=locals()
filename_up=locals()

#新頁面上傳操作
def uploadfunction(final_data):
    for index in final_data:
        titlename=index.get('title')
        splitname=titlename.split('.pdf')
        finalname=splitname[0]
        nowtime=datetime.datetime.now()
        # print(index.get('title'))
        # print(index.get('path'))

        Interface_FTP.ChooseFtpPath('/Uploads/outlook_doc/')    
        Interface_FTP.UploadToFtp(index.get('path'),index.get('title'))
        Interface_SQL.InsertToSQL(46,nowtime,finalname,'http://mis.nec.com.tw/uploads/outlook_doc/{}'.format(titlename),'更新時間: {}'.format(time.strftime("%Y.%m.%d %H:%M", time.localtime())))


def createNewWindow():
    newWindow = tk.Toplevel(window)
    newWindow.geometry("400x500+300+150")
    label_new=tk.Label(newWindow,text='選擇需要上傳至Dashboard的PDF檔')
    label_new.pack()
    current_path=os.getcwd()
    allfile=os.listdir(current_path)
    final_data=[]
    #upload_counter=len(allfile)
    for file_list in allfile:
        if '.pdf' in file_list:
            i=0
            upload_buttons_value['upload_buttons_value' + str(i) ] =  tk.BooleanVar()
            upload_buttons['upload_buttons' + str(i) ] = tk.Checkbutton(newWindow,variable=upload_buttons_value.get('upload_buttons_value' + str(i)),text=str(file_list).strip()) 
            #print(file_list)
            upload_buttons.get('upload_buttons' + str(i)).pack()
            filename_up['filename_up'+str(i)]=os.path.abspath(file_list.strip())
            #print(filename_up['filename_up'+str(i)])
            final_data.append({'title':file_list.strip(),'path':str(filename_up['filename_up'+str(i)])})
            i+=1
            print(final_data)
    if '.pdf' not in allfile:
        #print("fk")
        label_alert2 = tk.Label(newWindow, text = "資料夾中無PDF檔，請補上後再上傳!!!")
        label_alert2.pack()
            #print(file_list.strip())

    button_final = tk.Button(newWindow, text = "上傳至DashBoard",command=partial(uploadfunction,final_data)) #設置上傳function
    button_final.pack()
            


button_next = tk.Button(window, text="下一步",command=createNewWindow)
button_next.pack(side='bottom')

window.mainloop()